"""File processing utilities for different document formats."""

import os
import re
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from io import BytesIO

try:
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph
except ImportError:
    DocxDocument = None
    DocxTable = None
    Paragraph = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None

try:
    from PIL import Image
except ImportError:
    Image = None

from ..utils.logging_config import get_logger
from ..utils.error_handling import (
    create_error, ErrorCategory, ErrorSeverity, ProcessingError
)

logger = get_logger(__name__)


class FileProcessorError(ProcessingError):
    """Exception raised for file processing errors."""
    pass


class FileProcessor:
    """Handles processing of different file formats."""
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.md': 'text/markdown', 
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.pdf': 'application/pdf',
        '.rtf': 'application/rtf'
    }
    
    def __init__(self, show_extracted_text: bool = False):
        """Initialize file processor.
        
        Args:
            show_extracted_text: Whether to show extracted text during processing.
        """
        self.show_extracted_text = show_extracted_text
        
    def is_supported_format(self, file_path: Path) -> bool:
        """Check if file format is supported.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            True if format is supported, False otherwise.
        """
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions.
        
        Returns:
            List of supported extensions.
        """
        return list(self.SUPPORTED_EXTENSIONS.keys())
    
    def get_file_type_description(self, file_path: Path) -> str:
        """Get human-readable description of file type.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            File type description.
        """
        ext = file_path.suffix.lower()
        descriptions = {
            '.txt': 'Plain Text',
            '.md': 'Markdown',
            '.docx': 'Microsoft Word Document',
            '.pdf': 'PDF Document',
            '.rtf': 'Rich Text Format'
        }
        return descriptions.get(ext, f'Unknown ({ext})')
    
    def extract_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from file.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            Tuple of (extracted_text, metadata).
            
        Raises:
            FileProcessorError: If extraction fails.
        """
        if not file_path.exists():
            error = create_error(
                error_code="FILE_NOT_FOUND",
                message=f"File not found: {file_path}",
                category=ErrorCategory.VALIDATION,
                severity=ErrorSeverity.HIGH,
                details={'file_path': str(file_path)},
                suggestions=[
                    "Check if the file path is correct",
                    "Verify file permissions",
                    "Ensure the file exists"
                ]
            )
            raise FileProcessorError(error.error_info, error)
        
        if not self.is_supported_format(file_path):
            error = create_error(
                error_code="UNSUPPORTED_FILE_FORMAT",
                message=f"Unsupported file format: {file_path.suffix}",
                category=ErrorCategory.VALIDATION,
                severity=ErrorSeverity.MEDIUM,
                details={
                    'file_extension': file_path.suffix,
                    'supported_extensions': self.get_supported_extensions()
                },
                suggestions=[
                    f"Convert file to one of supported formats: {', '.join(self.get_supported_extensions())}",
                    "Check if required libraries are installed"
                ]
            )
            raise FileProcessorError(error.error_info, error)
        
        try:
            ext = file_path.suffix.lower()
            
            if ext == '.txt':
                return self._extract_text_file(file_path)
            elif ext == '.md':
                return self._extract_markdown_file(file_path)
            elif ext == '.docx':
                return self._extract_docx_file(file_path)
            elif ext == '.pdf':
                return self._extract_pdf_file(file_path)
            elif ext == '.rtf':
                return self._extract_rtf_file(file_path)
            else:
                raise FileProcessorError(f"Handler not implemented for {ext}")
                
        except FileProcessorError:
            raise
        except Exception as e:
            error = create_error(
                error_code="FILE_EXTRACTION_FAILED",
                message=f"Failed to extract text from {file_path.name}: {str(e)}",
                category=ErrorCategory.PROCESSING,
                severity=ErrorSeverity.HIGH,
                details={
                    'file_path': str(file_path),
                    'file_extension': file_path.suffix,
                    'error_details': str(e)
                },
                suggestions=[
                    "Check if the file is corrupted",
                    "Verify file encoding",
                    "Try opening the file manually",
                    "Check if required libraries are installed"
                ]
            )
            raise FileProcessorError(error.error_info, e)
    
    def _extract_text_file(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file.
        
        Args:
            file_path: Path to the text file.
            
        Returns:
            Tuple of (text_content, metadata).
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['cp1251', 'latin1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.warning(f"Used {encoding} encoding for {file_path}")
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise FileProcessorError(f"Could not decode file {file_path} with any encoding")
        
        # Clean and validate text
        content = self._clean_text(content)
        
        metadata = {
            'file_type': 'text',
            'encoding': 'utf-8',
            'character_count': len(content),
            'line_count': content.count('\n') + 1 if content else 0
        }
        
        if self.show_extracted_text:
            logger.info(f"Extracted text from {file_path.name}:\n{content[:500]}...")
        
        return content, metadata
    
    def _extract_markdown_file(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from markdown file.
        
        Args:
            file_path: Path to the markdown file.
            
        Returns:
            Tuple of (text_content, metadata).
        """
        # For now, treat markdown as plain text
        # In the future, we could parse markdown structure
        content, metadata = self._extract_text_file(file_path)
        metadata['file_type'] = 'markdown'
        
        # Count markdown elements
        headers = len(re.findall(r'^#+\s+', content, re.MULTILINE))
        links = len(re.findall(r'\[.*?\]\(.*?\)', content))
        code_blocks = len(re.findall(r'```.*?```', content, re.DOTALL))
        
        metadata.update({
            'headers_count': headers,
            'links_count': links,
            'code_blocks_count': code_blocks
        })
        
        if self.show_extracted_text:
            logger.info(f"Extracted markdown from {file_path.name}:\n{content[:500]}...")
        
        return content, metadata
    
    def _extract_docx_file(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file with enhanced table and structure support.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            Tuple of (text_content, metadata).
        """
        if DocxDocument is None:
            raise FileProcessorError("python-docx library not installed")
        
        try:
            doc = DocxDocument(file_path)
            
            content_parts = []
            tables_count = 0
            images_count = 0
            
            # Process document elements in order
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find corresponding paragraph object
                    for para in doc.paragraphs:
                        if para._element == element:
                            text = para.text.strip()
                            if text:
                                content_parts.append(text)
                            break
                
                elif element.tag.endswith('tbl'):  # Table
                    # Find corresponding table object
                    for table in doc.tables:
                        if table._element == element:
                            table_text = self._extract_table_text(table)
                            if table_text:
                                content_parts.append(f"\n[TABLE {tables_count + 1}]\n{table_text}\n[/TABLE]\n")
                                tables_count += 1
                            break
            
            # Fallback: if no content from element processing, use paragraphs directly
            if not content_parts:
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        content_parts.append(text)
            
            # Count images (drawings and inline shapes)
            for para in doc.paragraphs:
                images_count += len(para._element.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}))
            
            content = '\n'.join(content_parts)
            content = self._clean_text(content)
            
            metadata = {
                'file_type': 'docx',
                'character_count': len(content),
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                'tables_count': tables_count,
                'images_count': images_count,
                'sections_count': len(doc.sections)
            }
            
            # Add document properties if available
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata['document_title'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata['document_author'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                metadata['document_created'] = doc.core_properties.created.isoformat()
            
            if self.show_extracted_text:
                logger.info(f"Extracted DOCX from {file_path.name}:\n{content[:500]}...")
            
            return content, metadata
            
        except Exception as e:
            raise FileProcessorError(f"Failed to extract DOCX content: {str(e)}")
    
    def _extract_table_text(self, table: 'DocxTable') -> str:
        """Extract text from DOCX table with proper formatting.
        
        Args:
            table: DOCX table object.
            
        Returns:
            Formatted table text.
        """
        if not table or not table.rows:
            return ""
        
        table_lines = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = ' '.join(para.text.strip() for para in cell.paragraphs if para.text.strip())
                row_cells.append(cell_text or '')
            
            if any(cell.strip() for cell in row_cells):  # Skip empty rows
                table_lines.append(' | '.join(row_cells))
        
        return '\n'.join(table_lines)
    
    def _extract_pdf_file(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to the PDF file.
            
        Returns:
            Tuple of (text_content, metadata).
        """
        if PyPDF2 is None:
            raise FileProcessorError("PyPDF2 library not installed")
        
        try:
            content_parts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_parts.append(f"[PAGE {page_num + 1}]\n{page_text.strip()}\n")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                content = '\n'.join(content_parts)
                content = self._clean_text(content)
                
                metadata = {
                    'file_type': 'pdf',
                    'character_count': len(content),
                    'pages_count': len(pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted
                }
                
                # Add PDF metadata if available
                if pdf_reader.metadata:
                    pdf_info = pdf_reader.metadata
                    if pdf_info.get('/Title'):
                        metadata['document_title'] = str(pdf_info['/Title'])
                    if pdf_info.get('/Author'):
                        metadata['document_author'] = str(pdf_info['/Author'])
                    if pdf_info.get('/Creator'):
                        metadata['document_creator'] = str(pdf_info['/Creator'])
                    if pdf_info.get('/CreationDate'):
                        metadata['document_created'] = str(pdf_info['/CreationDate'])
                
                if self.show_extracted_text:
                    logger.info(f"Extracted PDF from {file_path.name}:\n{content[:500]}...")
                
                return content, metadata
                
        except Exception as e:
            raise FileProcessorError(f"Failed to extract PDF content: {str(e)}")
    
    def _extract_rtf_file(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from RTF file.
        
        Args:
            file_path: Path to the RTF file.
            
        Returns:
            Tuple of (text_content, metadata).
        """
        if rtf_to_text is None:
            raise FileProcessorError("striprtf library not installed")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['cp1251', 'latin1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        rtf_content = file.read()
                    logger.warning(f"Used {encoding} encoding for RTF file {file_path}")
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise FileProcessorError(f"Could not decode RTF file {file_path} with any encoding")
        
        try:
            # Convert RTF to plain text
            content = rtf_to_text(rtf_content)
            content = self._clean_text(content)
            
            metadata = {
                'file_type': 'rtf',
                'character_count': len(content),
                'original_rtf_size': len(rtf_content)
            }
            
            if self.show_extracted_text:
                logger.info(f"Extracted RTF from {file_path.name}:\n{content[:500]}...")
            
            return content, metadata
            
        except Exception as e:
            raise FileProcessorError(f"Failed to extract RTF content: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text.
            
        Returns:
            Cleaned text.
        """
        if not text:
            return ""
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize unicode characters
        text = text.encode('utf-8', errors='ignore').decode('utf-8')
        
        # Remove excessive whitespace but preserve paragraph structure
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        text = re.sub(r'^[ \t]+|[ \t]+$', '', text, flags=re.MULTILINE)  # Trim spaces/tabs from lines
        
        return text.strip()
    
    def validate_extracted_text(self, text: str, min_length: int = 10) -> bool:
        """Validate extracted text content.
        
        Args:
            text: Extracted text to validate.
            min_length: Minimum required text length.
            
        Returns:
            True if text is valid, False otherwise.
        """
        if not text or not text.strip():
            return False
        
        if len(text.strip()) < min_length:
            return False
        
        # Check if text contains mostly readable characters
        readable_chars = sum(1 for c in text if c.isprintable() or c.isspace())
        if readable_chars / len(text) < 0.8:  # At least 80% readable characters
            return False
        
        return True


# Global file processor instance
file_processor = FileProcessor()