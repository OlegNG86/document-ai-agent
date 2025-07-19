"""Document manager for handling document upload, processing and search."""

import os
import uuid
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import hashlib

import chromadb
from chromadb.config import Settings
from docx import Document as DocxDocument

from ..models.document import Document, DocumentCategory
from .ollama_client import OllamaClient, OllamaConnectionError


logger = logging.getLogger(__name__)


class DocumentManagerError(Exception):
    """Exception raised for document management errors."""
    pass


class DocumentManager:
    """Manages document storage, indexing and search using ChromaDB."""
    
    def __init__(
        self, 
        storage_path: str = "data/documents",
        chroma_path: str = "data/chroma_db",
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        """Initialize document manager.
        
        Args:
            storage_path: Path to store uploaded documents.
            chroma_path: Path for ChromaDB storage.
            chunk_size: Size of text chunks for vectorization.
            chunk_overlap: Overlap between chunks.
        """
        self.storage_path = Path(storage_path)
        self.chroma_path = Path(chroma_path)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Create directories if they don't exist
        self.storage_path.mkdir(parents=True, exist_ok=True)
        self.chroma_path.mkdir(parents=True, exist_ok=True)
        
        # Initialize ChromaDB
        self.chroma_client = chromadb.PersistentClient(
            path=str(self.chroma_path),
            settings=Settings(anonymized_telemetry=False)
        )
        
        # Get or create collection
        self.collection = self.chroma_client.get_or_create_collection(
            name="documents",
            metadata={"description": "Document knowledge base"}
        )
        
        # Initialize Ollama client for embeddings
        self.ollama_client = OllamaClient()
        
    def upload_document(
        self, 
        file_path: str, 
        metadata: Optional[Dict[str, Any]] = None,
        category: DocumentCategory = DocumentCategory.GENERAL,
        tags: Optional[List[str]] = None
    ) -> str:
        """Upload and process a document.
        
        Args:
            file_path: Path to the document file.
            metadata: Additional metadata for the document.
            category: Document category (reference, target, general).
            tags: List of tags for the document.
            
        Returns:
            Document ID.
            
        Raises:
            DocumentManagerError: If upload fails.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentManagerError(f"File not found: {file_path}")
        
        # Get supported file types from environment or use defaults
        supported_types = os.getenv('SUPPORTED_FILE_TYPES', '.txt,.md,.docx').split(',')
        supported_types = [ext.strip().lower() for ext in supported_types]
        
        if file_path.suffix.lower() not in supported_types:
            raise DocumentManagerError(f"Unsupported file type: {file_path.suffix}. Supported types: {', '.join(supported_types)}")
        
        try:
            # Read file content based on file type
            if file_path.suffix.lower() == '.docx':
                # Read DOCX file
                doc = DocxDocument(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            else:
                # Read text files (txt, md)
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            if not content.strip():
                raise DocumentManagerError("Document content is empty")
            
            # Generate document ID
            doc_id = str(uuid.uuid4())
            
            # Create document object
            document = Document(
                id=doc_id,
                title=file_path.stem,
                content=content,
                file_path=file_path,
                file_type=file_path.suffix[1:],  # Remove the dot
                category=category,
                tags=tags or [],
                metadata=metadata or {}
            )
            
            # Copy file to storage
            stored_file_path = self.storage_path / f"{doc_id}_{file_path.name}"
            with open(stored_file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # Update document file path to stored location
            document.file_path = stored_file_path
            
            # Split into chunks and create embeddings
            chunks = self._split_text_into_chunks(content)
            self._store_document_chunks(document, chunks)
            
            logger.info(f"Document uploaded successfully: {doc_id}")
            return doc_id
            
        except Exception as e:
            logger.error(f"Failed to upload document: {e}")
            raise DocumentManagerError(f"Upload failed: {e}")
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document and its chunks.
        
        Args:
            document_id: ID of document to delete.
            
        Returns:
            True if deleted successfully, False otherwise.
        """
        try:
            # Delete from ChromaDB
            chunk_ids = [f"{document_id}_chunk_{i}" for i in range(1000)]  # Max chunks assumption
            
            # Get existing chunk IDs for this document
            existing_chunks = self.collection.get(
                where={"document_id": document_id}
            )
            
            if existing_chunks['ids']:
                self.collection.delete(ids=existing_chunks['ids'])
            
            # Delete stored file
            for file_path in self.storage_path.glob(f"{document_id}_*"):
                file_path.unlink()
            
            logger.info(f"Document deleted: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document {document_id}: {e}")
            return False
    
    def get_document_info(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get information about a document.
        
        Args:
            document_id: ID of the document.
            
        Returns:
            Document information dict or None if not found.
        """
        try:
            # Get chunks for this document
            chunks_data = self.collection.get(
                where={"document_id": document_id}
            )
            
            if not chunks_data['ids']:
                return None
            
            # Get metadata from first chunk
            first_chunk_metadata = chunks_data['metadatas'][0]
            
            return {
                'id': document_id,
                'title': first_chunk_metadata.get('title', 'Unknown'),
                'file_type': first_chunk_metadata.get('file_type', 'unknown'),
                'category': first_chunk_metadata.get('category', 'general'),
                'tags': first_chunk_metadata.get('tags', '').split(',') if first_chunk_metadata.get('tags') else [],
                'upload_date': first_chunk_metadata.get('upload_date'),
                'chunk_count': len(chunks_data['ids']),
                'metadata': {k: v for k, v in first_chunk_metadata.items() 
                           if k not in ['document_id', 'chunk_index', 'title', 'file_type', 'upload_date', 'category', 'tags']}
            }
            
        except Exception as e:
            logger.error(f"Failed to get document info for {document_id}: {e}")
            return None
    
    def list_documents(
        self, 
        category_filter: Optional[DocumentCategory] = None,
        tags_filter: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """List all documents in the knowledge base.
        
        Args:
            category_filter: Filter by document category.
            tags_filter: Filter by document tags.
        
        Returns:
            List of document information dicts.
        """
        try:
            # Build where clause for filtering
            where_clause = {}
            if category_filter:
                where_clause['category'] = category_filter.value
            
            # Get chunks with optional filtering
            if where_clause:
                all_chunks = self.collection.get(where=where_clause)
            else:
                all_chunks = self.collection.get()
            
            # Group by document_id
            documents = {}
            for i, chunk_id in enumerate(all_chunks['ids']):
                metadata = all_chunks['metadatas'][i]
                doc_id = metadata['document_id']
                
                # Apply tags filter if specified
                if tags_filter:
                    doc_tags = metadata.get('tags', [])
                    if isinstance(doc_tags, str):
                        doc_tags = doc_tags.split(',') if doc_tags else []
                    if not any(tag in doc_tags for tag in tags_filter):
                        continue
                
                if doc_id not in documents:
                    documents[doc_id] = {
                        'id': doc_id,
                        'title': metadata.get('title', 'Unknown'),
                        'file_type': metadata.get('file_type', 'unknown'),
                        'category': metadata.get('category', 'general'),
                        'tags': metadata.get('tags', []),
                        'upload_date': metadata.get('upload_date'),
                        'chunk_count': 0,
                        'metadata': {k: v for k, v in metadata.items() 
                                   if k not in ['document_id', 'chunk_index', 'title', 'file_type', 'upload_date', 'category', 'tags']}
                    }
                
                documents[doc_id]['chunk_count'] += 1
            
            return list(documents.values())
            
        except Exception as e:
            logger.error(f"Failed to list documents: {e}")
            return []
    
    def search_similar_chunks(
        self, 
        query: str, 
        top_k: int = 5,
        category_filter: Optional[DocumentCategory] = None,
        tags_filter: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """Search for similar document chunks.
        
        Args:
            query: Search query.
            top_k: Number of top results to return.
            category_filter: Filter by document category.
            tags_filter: Filter by document tags.
            
        Returns:
            List of similar chunks with metadata.
            
        Raises:
            DocumentManagerError: If search fails.
        """
        try:
            # Generate query embedding using Ollama
            query_embedding = self.ollama_client.generate_embeddings(query)
            
            # Build where clause for filtering
            where_clause = {}
            if category_filter:
                where_clause['category'] = category_filter.value
            
            # Search in ChromaDB
            search_params = {
                'query_embeddings': [query_embedding],
                'n_results': top_k * 2,  # Get more results for filtering
                'include': ['documents', 'metadatas', 'distances']
            }
            
            if where_clause:
                search_params['where'] = where_clause
            
            results = self.collection.query(**search_params)
            
            # Format and filter results
            similar_chunks = []
            for i in range(len(results['documents'][0])):
                metadata = results['metadatas'][0][i]
                
                # Apply tags filter if specified
                if tags_filter:
                    doc_tags = metadata.get('tags', [])
                    if isinstance(doc_tags, str):
                        doc_tags = doc_tags.split(',') if doc_tags else []
                    if not any(tag in doc_tags for tag in tags_filter):
                        continue
                
                chunk_data = {
                    'id': results['ids'][0][i],
                    'content': results['documents'][0][i],
                    'metadata': metadata,
                    'distance': results['distances'][0][i],
                    'relevance_score': max(0.0, min(1.0, 1.0 - results['distances'][0][i]))  # Convert distance to similarity, clamp to [0,1]
                }
                similar_chunks.append(chunk_data)
                
                # Stop when we have enough results
                if len(similar_chunks) >= top_k:
                    break
            
            return similar_chunks
            
        except OllamaConnectionError as e:
            logger.error(f"Ollama connection error during search: {e}")
            raise DocumentManagerError(f"Search failed - Ollama unavailable: {e}")
        except Exception as e:
            logger.error(f"Search failed: {e}")
            raise DocumentManagerError(f"Search failed: {e}")
    
    def _split_text_into_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks.
        
        Args:
            text: Text to split.
            
        Returns:
            List of text chunks.
        """
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Find the last sentence ending before the chunk limit
            chunk_text = text[start:end]
            last_sentence_end = max(
                chunk_text.rfind('.'),
                chunk_text.rfind('!'),
                chunk_text.rfind('?'),
                chunk_text.rfind('\n')
            )
            
            if last_sentence_end > self.chunk_size // 2:  # Ensure chunk isn't too small
                end = start + last_sentence_end + 1
            
            chunks.append(text[start:end])
            start = end - self.chunk_overlap
        
        return chunks
    
    def _store_document_chunks(self, document: Document, chunks: List[str]) -> None:
        """Store document chunks in ChromaDB.
        
        Args:
            document: Document object.
            chunks: List of text chunks.
            
        Raises:
            DocumentManagerError: If storage fails.
        """
        try:
            chunk_ids = []
            chunk_embeddings = []
            chunk_documents = []
            chunk_metadatas = []
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"{document.id}_chunk_{i}"
                
                # Generate embedding for chunk
                embedding = self.ollama_client.generate_embeddings(chunk)
                
                # Prepare metadata
                metadata = {
                    'document_id': document.id,
                    'chunk_index': i,
                    'title': document.title,
                    'file_type': document.file_type,
                    'category': document.category.value,
                    'tags': ','.join(document.tags) if document.tags else '',
                    'upload_date': document.created_at.isoformat(),
                    **document.metadata
                }
                
                chunk_ids.append(chunk_id)
                chunk_embeddings.append(embedding)
                chunk_documents.append(chunk)
                chunk_metadatas.append(metadata)
                
                # Update document object
                document.add_chunk(chunk, chunk_id)
            
            # Store in ChromaDB
            self.collection.add(
                ids=chunk_ids,
                embeddings=chunk_embeddings,
                documents=chunk_documents,
                metadatas=chunk_metadatas
            )
            
            logger.info(f"Stored {len(chunks)} chunks for document {document.id}")
            
        except OllamaConnectionError as e:
            logger.error(f"Ollama connection error during chunk storage: {e}")
            raise DocumentManagerError(f"Failed to generate embeddings: {e}")
        except Exception as e:
            logger.error(f"Failed to store document chunks: {e}")
            raise DocumentManagerError(f"Chunk storage failed: {e}")
    
    def batch_upload_documents(
        self, 
        file_paths: List[str], 
        common_metadata: Optional[Dict[str, Any]] = None,
        category: DocumentCategory = DocumentCategory.GENERAL,
        tags: Optional[List[str]] = None,
        progress_callback: Optional[callable] = None
    ) -> List[Dict[str, Any]]:
        """Upload multiple documents in batch.
        
        Args:
            file_paths: List of file paths to upload.
            common_metadata: Common metadata to apply to all documents.
            category: Category to apply to all documents.
            tags: Tags to apply to all documents.
            progress_callback: Optional callback function for progress updates.
            
        Returns:
            List of results for each file upload.
        """
        results = []
        common_metadata = common_metadata or {}
        
        for i, file_path in enumerate(file_paths):
            try:
                # Prepare metadata for this file
                file_metadata = common_metadata.copy()
                file_metadata['batch_upload'] = 'true'
                file_metadata['batch_index'] = str(i)
                
                # Upload document
                doc_id = self.upload_document(
                    file_path=file_path, 
                    metadata=file_metadata,
                    category=category,
                    tags=tags
                )
                
                result = {
                    'success': True,
                    'file_path': file_path,
                    'document_id': doc_id,
                    'error': None
                }
                
            except Exception as e:
                result = {
                    'success': False,
                    'file_path': file_path,
                    'document_id': None,
                    'error': str(e)
                }
                logger.error(f"Failed to upload {file_path} in batch: {e}")
            
            results.append(result)
            
            # Call progress callback if provided
            if progress_callback:
                progress_callback(i + 1, len(file_paths), file_path, result)
        
        return results
    
    def update_document_category(self, document_id: str, category: DocumentCategory) -> bool:
        """Update document category.
        
        Args:
            document_id: ID of document to update.
            category: New category for the document.
            
        Returns:
            True if updated successfully, False otherwise.
        """
        try:
            # Get all chunks for this document
            chunks_data = self.collection.get(
                where={"document_id": document_id}
            )
            
            if not chunks_data['ids']:
                return False
            
            # Update metadata for all chunks
            updated_metadatas = []
            for metadata in chunks_data['metadatas']:
                metadata['category'] = category.value
                updated_metadatas.append(metadata)
            
            # Update in ChromaDB
            self.collection.update(
                ids=chunks_data['ids'],
                metadatas=updated_metadatas
            )
            
            logger.info(f"Updated category for document {document_id} to {category.value}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to update document category {document_id}: {e}")
            return False
    
    def update_document_tags(self, document_id: str, tags: List[str]) -> bool:
        """Update document tags.
        
        Args:
            document_id: ID of document to update.
            tags: New tags for the document.
            
        Returns:
            True if updated successfully, False otherwise.
        """
        try:
            # Get all chunks for this document
            chunks_data = self.collection.get(
                where={"document_id": document_id}
            )
            
            if not chunks_data['ids']:
                return False
            
            # Update metadata for all chunks
            updated_metadatas = []
            for metadata in chunks_data['metadatas']:
                metadata['tags'] = ','.join(tags) if tags else ''
                updated_metadatas.append(metadata)
            
            # Update in ChromaDB
            self.collection.update(
                ids=chunks_data['ids'],
                metadatas=updated_metadatas
            )
            
            logger.info(f"Updated tags for document {document_id} to {tags}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to update document tags {document_id}: {e}")
            return False
    
    def get_reference_documents(self) -> List[Dict[str, Any]]:
        """Get all reference/normative documents.
        
        Returns:
            List of reference document information.
        """
        return self.list_documents(category_filter=DocumentCategory.REFERENCE)
    
    def search_reference_chunks(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for similar chunks only in reference documents.
        
        Args:
            query: Search query.
            top_k: Number of top results to return.
            
        Returns:
            List of similar chunks from reference documents.
        """
        return self.search_similar_chunks(
            query=query,
            top_k=top_k,
            category_filter=DocumentCategory.REFERENCE
        )
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the document collection.
        
        Returns:
            Collection statistics.
        """
        try:
            collection_data = self.collection.get()
            total_chunks = len(collection_data['ids'])
            
            # Count unique documents and categories
            document_ids = set()
            category_counts = {}
            
            for metadata in collection_data['metadatas']:
                document_ids.add(metadata['document_id'])
                category = metadata.get('category', 'general')
                category_counts[category] = category_counts.get(category, 0) + 1
            
            # Convert chunk counts to document counts
            documents_by_category = {}
            processed_docs = set()
            
            for metadata in collection_data['metadatas']:
                doc_id = metadata['document_id']
                if doc_id not in processed_docs:
                    category = metadata.get('category', 'general')
                    documents_by_category[category] = documents_by_category.get(category, 0) + 1
                    processed_docs.add(doc_id)
            
            return {
                'total_documents': len(document_ids),
                'total_chunks': total_chunks,
                'average_chunks_per_document': total_chunks / len(document_ids) if document_ids else 0,
                'documents_by_category': documents_by_category,
                'chunks_by_category': category_counts
            }
            
        except Exception as e:
            logger.error(f"Failed to get collection stats: {e}")
            return {'error': str(e)}